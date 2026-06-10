"""文档解析与清洗：PDF / MD / TXT / DOCX。"""

from __future__ import annotations

import logging
import re
import urllib.error
import urllib.request
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_BODY_CHARS = 80_000
EXCERPT_CHARS = 200

_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_MD_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_MD_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
_MD_HEADING_RE = re.compile(r"^#{1,6}\s+")


@dataclass
class ParsedDocument:
    title: str
    body: str
    excerpt: str
    source_filename: str
    mime: str
    char_count: int


def _strip_markdown_inline(text: str) -> str:
    """标题等字段：去掉 **加粗**、*斜体*、`代码` 标记，保留可读文字。"""
    if not text:
        return ""
    out = _MD_INLINE_CODE_RE.sub(r"\1", text)
    out = _MD_BOLD_RE.sub(r"\1", out)
    out = _MD_ITALIC_RE.sub(r"\1", out)
    return out.strip()


def _plain_excerpt(body: str) -> str:
    plain = _strip_markdown_inline(body)
    plain = _MD_HEADING_RE.sub("", plain)
    plain = re.sub(r"\s+", " ", plain).strip()
    if len(plain) <= EXCERPT_CHARS:
        return plain
    return plain[:EXCERPT_CHARS] + "…"


def _drop_leading_title_line(body: str, title: str) -> str:
    """正文首行若为与标题一致的 # 标题，去掉避免发布重复。"""
    lines = body.splitlines()
    if not lines:
        return body
    first = lines[0].strip()
    m = re.match(r"^#{1,6}\s+(.+)$", first)
    if not m:
        return body
    heading = _strip_markdown_inline(m.group(1).strip())
    title_clean = _strip_markdown_inline(title.strip())
    if heading == title_clean:
        rest = "\n".join(lines[1:]).lstrip("\n")
        return rest if rest else body
    return body


def _clean_text(text: str) -> str:
    if not text:
        return ""
    out = text.replace("\ufeff", "")
    out = out.replace("\r\n", "\n").replace("\r", "\n")
    out = re.sub(r"[ \t]+\n", "\n", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def _title_from_markdown(body: str, fallback: str) -> str:
    for line in body.splitlines():
        m = re.match(r"^#\s+(.+)$", line.strip())
        if m:
            return _strip_markdown_inline(m.group(1).strip())[:255]
    return _strip_markdown_inline(_title_from_filename(fallback))[:255]


def _title_from_filename(name: str) -> str:
    stem = Path(name).stem.strip()
    return stem[:255] if stem else "未命名笔记"


def _truncate(body: str) -> str:
    if len(body) <= MAX_BODY_CHARS:
        return body
    return body[:MAX_BODY_CHARS] + "\n\n…（内容已截断）"


def _resolve_fetch_url(url: str) -> str:
    raw = (url or "").strip()
    if not raw:
        raise ValueError("empty url")
    if raw.startswith("http://") or raw.startswith("https://"):
        return raw
    from utils.qq.qq_music_tools import blog_api_base

    base = blog_api_base().rstrip("/")
    if raw.startswith("/v1/"):
        return f"{base}{raw[3:]}"
    if raw.startswith("/"):
        return f"{base}{raw}"
    return f"{base}/{raw.lstrip('/')}"


def _fetch_bytes(url: str, *, access_token: str = "") -> bytes:
    fetch_url = _resolve_fetch_url(url)
    headers = {"User-Agent": "blog-agent/1.0"}
    token = (access_token or "").strip()
    if token:
        headers["Authorization"] = f"Bearer {token}"
    req = urllib.request.Request(fetch_url, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            return resp.read()
    except urllib.error.HTTPError as exc:
        logger.warning(
            "[document] fetch failed url=%s resolved=%s status=%s reason=%s",
            url,
            fetch_url,
            exc.code,
            exc.reason,
        )
        raise


def parse_text_content(raw: str, *, filename: str, mime: str = "text/plain") -> ParsedDocument:
    body = _truncate(_clean_text(raw))
    title = _title_from_markdown(body, _title_from_filename(filename))
    body = _drop_leading_title_line(body, title)
    excerpt = _plain_excerpt(body)
    return ParsedDocument(
        title=title,
        body=body,
        excerpt=excerpt,
        source_filename=filename,
        mime=mime,
        char_count=len(body),
    )


def parse_pdf_bytes(data: bytes, *, filename: str) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("未安装 pypdf，无法解析 PDF") from exc
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            piece = page.extract_text() or ""
        except Exception:
            piece = ""
        if piece.strip():
            parts.append(piece)
    body = _truncate(_clean_text("\n\n".join(parts)))
    title = _title_from_filename(filename)
    if body:
        first = body.split("\n", 1)[0].strip()
        if 4 <= len(first) <= 80:
            title = first
    excerpt = body[:EXCERPT_CHARS] + ("…" if len(body) > EXCERPT_CHARS else "")
    return ParsedDocument(
        title=title,
        body=body,
        excerpt=excerpt,
        source_filename=filename,
        mime="application/pdf",
        char_count=len(body),
    )


def parse_docx_bytes(data: bytes, *, filename: str) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("未安装 python-docx，无法解析 DOCX") from exc
    doc = Document(BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    body = _truncate(_clean_text("\n\n".join(parts)))
    title = _title_from_filename(filename)
    excerpt = body[:EXCERPT_CHARS] + ("…" if len(body) > EXCERPT_CHARS else "")
    return ParsedDocument(
        title=title,
        body=body,
        excerpt=excerpt,
        source_filename=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        char_count=len(body),
    )


def parse_attachment(*, url: str, name: str, mime: str = "", kind: str = "") -> ParsedDocument:
    filename = (name or "document").strip() or "document"
    lower = filename.lower()
    mime_l = (mime or "").lower()
    kind_l = (kind or "").lower()

    # 有 URL 时一律拉取解析；内联 text 由 publish_note_route 走 parse_text_content
    if kind_l == "image" or mime_l.startswith("image/"):
        raise ValueError(f"图片附件不支持文档解析：{filename}")

    try:
        data = _fetch_bytes(url)
    except urllib.error.URLError as exc:
        logger.warning("[document] fetch failed url=%s err=%s", url, exc.reason)
        raise RuntimeError(f"无法读取附件：{exc.reason}") from exc
    except urllib.error.HTTPError as exc:
        raise RuntimeError(f"无法读取附件：HTTP {exc.code}") from exc

    if lower.endswith(".pdf") or mime_l == "application/pdf":
        return parse_pdf_bytes(data, filename=filename)
    if lower.endswith(".docx") or "wordprocessingml" in mime_l:
        return parse_docx_bytes(data, filename=filename)
    if lower.endswith((".md", ".markdown", ".txt")) or mime_l in ("text/markdown", "text/plain"):
        raw = data.decode("utf-8", errors="replace")
        return parse_text_content(raw, filename=filename, mime=mime or "text/plain")

    raise ValueError(f"暂不支持的文档类型：{filename}")


def merge_parsed_documents(docs: list[ParsedDocument]) -> ParsedDocument:
    if not docs:
        return ParsedDocument("", "", "", "", "", 0)
    if len(docs) == 1:
        return docs[0]
    title = docs[0].title
    names = ", ".join(d.source_filename for d in docs)
    body = _truncate(
        _clean_text("\n\n---\n\n".join(f"## {d.source_filename}\n\n{d.body}" for d in docs))
    )
    excerpt = _plain_excerpt(body)
    return ParsedDocument(
        title=title,
        body=body,
        excerpt=excerpt,
        source_filename=names,
        mime="multipart/document",
        char_count=len(body),
    )
